
import os
import pandas as pd
from docx import Document
from datetime import datetime, timedelta
import random

def create_sample_txt():
    
    sample_text = """
    Artificial Intelligence and Machine Learning: A Comprehensive Overview
    
    Artificial Intelligence (AI) and Machine Learning (ML) have revolutionized the way we approach problem-solving and data analysis. These technologies have become integral parts of modern computing and are driving innovation across various industries.
    
    Machine Learning, a subset of AI, enables computers to learn and improve from experience without being explicitly programmed. This is achieved through algorithms that can identify patterns in data and make predictions or decisions based on those patterns.
    
    There are three main types of machine learning: supervised learning, unsupervised learning, and reinforcement learning. Supervised learning involves training a model on labeled data, while unsupervised learning finds hidden patterns in unlabeled data. Reinforcement learning uses a system of rewards and penalties to guide the learning process.
    
    Deep Learning, a subset of machine learning, uses neural networks with multiple layers to process complex data. This approach has been particularly successful in image recognition, natural language processing, and speech recognition.
    
    The applications of AI and ML are vast and growing. They are used in healthcare for disease diagnosis, in finance for fraud detection, in transportation for autonomous vehicles, and in many other fields. The potential for these technologies to improve efficiency and create new opportunities is enormous.
    
    However, the development and deployment of AI systems also raise important ethical considerations. Issues such as bias in algorithms, privacy concerns, and the impact on employment need to be carefully considered and addressed.
    
    As we move forward, it is crucial to develop AI systems that are not only effective but also fair, transparent, and beneficial to society as a whole. This requires collaboration between technologists, policymakers, and other stakeholders.
    """
    
    with open("sample_document.txt", "w", encoding="utf-8") as f:
        f.write(sample_text.strip())
    
    print(" Created sample_document.txt")

def create_sample_csv():
    dates = pd.date_range(start='2023-01-01', end='2023-12-31', freq='D')
    products = ['Laptop', 'Smartphone', 'Tablet', 'Headphones', 'Monitor']
    regions = ['North', 'South', 'East', 'West']
    
    data = []
    for date in dates:
        for _ in range(random.randint(5, 15)):  
            data.append({
                'Date': date.strftime('%Y-%m-%d'),
                'Product': random.choice(products),
                'Region': random.choice(regions),
                'Quantity': random.randint(1, 10),
                'Unit_Price': round(random.uniform(100, 2000), 2),
                'Sales_Rep': f"Rep_{random.randint(1, 20)}"
            })
    
    df = pd.DataFrame(data)
    df['Total_Revenue'] = df['Quantity'] * df['Unit_Price']
    
    df.to_csv("sample_sales_data.csv", index=False)
    print("Created sample_sales_data.csv")

def create_sample_excel():
    
    with pd.ExcelWriter("sample_company_data.xlsx", engine='openpyxl') as writer:
        
        
        employees = []
        departments = ['Engineering', 'Sales', 'Marketing', 'HR', 'Finance']
        positions = ['Manager', 'Senior', 'Junior', 'Intern']
        
        for i in range(50):
            employees.append({
                'Employee_ID': f"EMP{i+1:03d}",
                'Name': f"Employee {i+1}",
                'Department': random.choice(departments),
                'Position': random.choice(positions),
                'Salary': random.randint(30000, 120000),
                'Hire_Date': (datetime.now() - timedelta(days=random.randint(30, 1000))).strftime('%Y-%m-%d'),
                'Performance_Rating': round(random.uniform(3.0, 5.0), 1)
            })
        
        df_employees = pd.DataFrame(employees)
        df_employees.to_excel(writer, sheet_name='Employees', index=False)
        
        
        projects = []
        statuses = ['Active', 'Completed', 'On Hold', 'Cancelled']
        
        for i in range(20):
            projects.append({
                'Project_ID': f"PROJ{i+1:03d}",
                'Project_Name': f"Project {i+1}",
                'Department': random.choice(departments),
                'Status': random.choice(statuses),
                'Budget': random.randint(50000, 500000),
                'Start_Date': (datetime.now() - timedelta(days=random.randint(100, 500))).strftime('%Y-%m-%d'),
                'End_Date': (datetime.now() + timedelta(days=random.randint(50, 300))).strftime('%Y-%m-%d'),
                'Team_Size': random.randint(3, 15)
            })
        
        df_projects = pd.DataFrame(projects)
        df_projects.to_excel(writer, sheet_name='Projects', index=False)
        
        
        months = pd.date_range(start='2023-01-01', end='2023-12-31', freq='M')
        financial_data = []
        
        for month in months:
            financial_data.append({
                'Month': month.strftime('%Y-%m'),
                'Revenue': random.randint(500000, 2000000),
                'Expenses': random.randint(300000, 1500000),
                'Profit': random.randint(100000, 800000),
                'New_Customers': random.randint(50, 200),
                'Customer_Satisfaction': round(random.uniform(3.5, 5.0), 2)
            })
        
        df_financial = pd.DataFrame(financial_data)
        df_financial.to_excel(writer, sheet_name='Financial_Summary', index=False)
    
    print("Created sample_company_data.xlsx")

def create_sample_docx():
    
    doc = Document()
    
   
    title = doc.add_heading('Business Strategy Report 2024', 0)
    
    
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph('''
    This report outlines our comprehensive business strategy for 2024, focusing on digital transformation, 
    market expansion, and operational efficiency. Our analysis indicates significant opportunities for 
    growth in emerging markets and technology sectors.
    ''')
    
   
    doc.add_heading('Market Analysis', level=1)
    doc.add_paragraph('''
    The current market landscape presents both challenges and opportunities. Our research shows that 
    consumer preferences are shifting towards digital solutions, creating new revenue streams for 
    companies that can adapt quickly.
    ''')
    
    
    doc.add_heading('Key Performance Indicators', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = '2023'
    hdr_cells[2].text = '2024 Target'
    
   
    metrics = [
        ('Revenue Growth', '15%', '25%'),
        ('Customer Acquisition', '1,200', '2,000'),
        ('Market Share', '8%', '12%'),
        ('Employee Satisfaction', '4.2/5', '4.5/5')
    ]
    
    for metric, current, target in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = current
        row_cells[2].text = target
    
    
    doc.add_heading('Strategic Initiatives', level=1)
    doc.add_paragraph('''
    Our strategic initiatives for 2024 include:
    ''')
    
    initiatives = [
        'Digital Transformation: Implement AI and automation across all business processes',
        'Market Expansion: Enter three new international markets',
        'Product Innovation: Launch five new products in the technology sector',
        'Talent Development: Invest in employee training and development programs',
        'Sustainability: Achieve carbon neutrality by 2025'
    ]
    
    for initiative in initiatives:
        doc.add_paragraph(initiative, style='List Bullet')
    
    
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph('''
    The proposed strategy positions our company for sustainable growth and market leadership. 
    Success will depend on effective execution, continuous monitoring, and adaptability to 
    changing market conditions.
    ''')
    
    doc.save("sample_business_report.docx")
    print("Created sample_business_report.docx")

def create_sample_pdf():
   
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        doc = SimpleDocTemplate("sample_research_paper.pdf", pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  
        )
        title = Paragraph("Research Paper: Climate Change Impact Analysis", title_style)
        story.append(title)
        story.append(Spacer(1, 12))
        
        
        abstract_style = ParagraphStyle(
            'Abstract',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=12,
            leftIndent=20,
            rightIndent=20
        )
        
        abstract_text = """
        This research paper examines the multifaceted impacts of climate change on global ecosystems, 
        economies, and human societies. Through comprehensive analysis of historical data and 
        predictive modeling, we identify key trends and potential mitigation strategies. Our findings 
        suggest that immediate action is required to address the most severe consequences of climate change.
        """
        
        story.append(Paragraph("Abstract", styles['Heading2']))
        story.append(Paragraph(abstract_text, abstract_style))
        story.append(Spacer(1, 12))
        
        
        intro_text = """
        Climate change represents one of the most significant challenges facing humanity in the 21st century. 
        The scientific consensus is clear: human activities, particularly the burning of fossil fuels, 
        have led to unprecedented increases in greenhouse gas concentrations in the atmosphere. This paper 
        provides a comprehensive analysis of the current state of climate change research and its implications 
        for policy and action.
        """
        
        story.append(Paragraph("Introduction", styles['Heading2']))
        story.append(Paragraph(intro_text, styles['Normal']))
        story.append(Spacer(1, 12))
        
        
        method_text = """
        Our research methodology combines quantitative analysis of climate data with qualitative assessment 
        of policy responses. We analyzed temperature records from 1950 to 2023, precipitation patterns, 
        and sea level measurements from multiple global monitoring stations. Additionally, we reviewed 
        over 200 peer-reviewed studies on climate change impacts and mitigation strategies.
        """
        
        story.append(Paragraph("Methodology", styles['Heading2']))
        story.append(Paragraph(method_text, styles['Normal']))
        story.append(Spacer(1, 12))
        
        results_text = """
        The analysis reveals several concerning trends. Global average temperatures have increased by 
        1.1°C since pre-industrial times, with the rate of warming accelerating in recent decades. 
        Extreme weather events have become more frequent and intense, affecting millions of people 
        worldwide. Sea levels are rising at an average rate of 3.3mm per year, threatening coastal 
        communities and ecosystems.
        """
        
        story.append(Paragraph("Results", styles['Heading2']))
        story.append(Paragraph(results_text, styles['Normal']))
        story.append(Spacer(1, 12))
        
        
        conclusion_text = """
        The evidence presented in this paper underscores the urgent need for comprehensive climate action. 
        While the challenges are significant, there are also opportunities for innovation and sustainable 
        development. Success will require unprecedented cooperation between nations, industries, and 
        individuals to reduce emissions and adapt to changing conditions.
        """
        
        story.append(Paragraph("Conclusion", styles['Heading2']))
        story.append(Paragraph(conclusion_text, styles['Normal']))
        
        doc.build(story)
        print(" Created sample_research_paper.pdf")
        
    except ImportError:
        print("ReportLab not installed. Skipping PDF creation.")
        print("   Install with: pip install reportlab")

def main():
   
    print(" Generating sample documents for testing...")
    print()
    
    
    if not os.path.exists("sample_files"):
        os.makedirs("sample_files")
    
    
    original_dir = os.getcwd()
    os.chdir("sample_files")
    
    try:
        create_sample_txt()
        create_sample_csv()
        create_sample_excel()
        create_sample_docx()
        create_sample_pdf()
        
        print()
        print("All sample files created successfully!")
        print("Files are located in the 'sample_files' directory")
        print()
        print(" Sample files created:")
        print("   • sample_document.txt - Text document about AI/ML")
        print("   • sample_sales_data.csv - Sales data with 1000+ records")
        print("   • sample_company_data.xlsx - Multi-sheet company data")
        print("   • sample_business_report.docx - Business strategy report")
        print("   • sample_research_paper.pdf - Research paper on climate change")
        print()
        print(" You can now upload these files to test the Document Summarizer!")
        
    except Exception as e:
        print(f" Error creating sample files: {str(e)}")
    finally:
        os.chdir(original_dir)

if __name__ == "__main__":
    main() 