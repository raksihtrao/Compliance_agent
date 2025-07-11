import fitz  # PyMuPDF
import pandas as pd
from docx import Document
import io
import re

class DocumentProcessor:
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._extract_pdf_text,
            'text/plain': self._extract_txt_text,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._extract_excel_text,
            'application/vnd.ms-excel': self._extract_excel_text,
            'text/csv': self._extract_csv_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'application/msword': self._extract_docx_text
        }
    
    def extract_text(self, uploaded_file):
       
        try:
           
            file_type = uploaded_file.type
            
            
            if file_type in self.supported_formats:
                return self.supported_formats[file_type](uploaded_file)
            else:
                
                file_extension = uploaded_file.name.lower().split('.')[-1]
                return self._extract_by_extension(uploaded_file, file_extension)
                
        except Exception as e:
            raise Exception(f"Error extracting text from {uploaded_file.name}: {str(e)}")
    
    def _extract_by_extension(self, uploaded_file, extension):
        
        extension_mapping = {
            'pdf': self._extract_pdf_text,
            'txt': self._extract_txt_text,
            'xlsx': self._extract_excel_text,
            'xls': self._extract_excel_text,
            'csv': self._extract_csv_text,
            'docx': self._extract_docx_text,
            'doc': self._extract_docx_text
        }
        
        if extension in extension_mapping:
            return extension_mapping[extension](uploaded_file)
        else:
            raise Exception(f"Unsupported file extension: {extension}")
    
    def _extract_pdf_text(self, uploaded_file):
       
        try:
            
            pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            text = ""
            
            
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text += page.get_text()
            
            pdf_document.close()
            
            
            text = self._clean_text(text)
            return text
            
        except Exception as e:
            raise Exception(f"Error extracting PDF text: {str(e)}")
    
    def _extract_txt_text(self, uploaded_file):
        
        try:
          
            text = uploaded_file.read().decode('utf-8')
            
           
            text = self._clean_text(text)
            return text
            
        except UnicodeDecodeError:
           
            try:
                uploaded_file.seek(0) 
                text = uploaded_file.read().decode('latin-1')
                text = self._clean_text(text)
                return text
            except Exception as e:
                raise Exception(f"Error decoding text file: {str(e)}")
        except Exception as e:
            raise Exception(f"Error extracting text file: {str(e)}")
    
    def _extract_excel_text(self, uploaded_file):
        
        try:
           
            df_list = pd.read_excel(uploaded_file, sheet_name=None)
            
            text_parts = []
            
            
            for sheet_name, df in df_list.items():
                text_parts.append(f"Sheet: {sheet_name}")
                
               
                if not df.empty:
                   
                    text_parts.append("Columns: " + ", ".join(df.columns.astype(str)))
                    
                    
                    text_parts.append("Data:")
                    text_parts.append(df.head(20).to_string(index=False))
                    
                    
                    numeric_cols = df.select_dtypes(include=['number']).columns
                    if len(numeric_cols) > 0:
                        text_parts.append("Summary Statistics:")
                        text_parts.append(df[numeric_cols].describe().to_string())
                
                text_parts.append("\n" + "="*50 + "\n")
            
            text = "\n".join(text_parts)
            text = self._clean_text(text)
            return text
            
        except Exception as e:
            raise Exception(f"Error extracting Excel text: {str(e)}")
    
    def _extract_csv_text(self, uploaded_file):
       
        try:
           
            df = pd.read_csv(uploaded_file)
            
            text_parts = []
            
            
            text_parts.append("Columns: " + ", ".join(df.columns.astype(str)))
            
           
            if not df.empty:
                text_parts.append("Data Preview:")
                text_parts.append(df.head(20).to_string(index=False))
                
                
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    text_parts.append("Summary Statistics:")
                    text_parts.append(df[numeric_cols].describe().to_string())
            
            text = "\n".join(text_parts)
            text = self._clean_text(text)
            return text
            
        except Exception as e:
            raise Exception(f"Error extracting CSV text: {str(e)}")
    
    def _extract_docx_text(self, uploaded_file):
        
        try:
            
            doc = Document(uploaded_file)
            
            text_parts = []
            
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
        
            for table in doc.tables:
                text_parts.append("\nTable:")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            text = self._clean_text(text)
            return text
            
        except Exception as e:
            raise Exception(f"Error extracting Word document text: {str(e)}")
    
    def _clean_text(self, text):
        
        if not text:
            return ""
        
        text = re.sub(r'\s+', ' ', text)
        
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        text = text.strip()
        
        return text
    
    def get_file_info(self, uploaded_file):
        try:
            file_size = len(uploaded_file.getvalue())
            file_type = uploaded_file.type
            file_name = uploaded_file.name
            
            return {
                'name': file_name,
                'type': file_type,
                'size_bytes': file_size,
                'size_kb': file_size / 1024,
                'size_mb': file_size / (1024 * 1024)
            }
        except Exception as e:
            raise Exception(f"Error getting file info: {str(e)}") 