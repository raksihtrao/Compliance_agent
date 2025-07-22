import os
import sys
import tempfile
import unittest
from unittest.mock import Mock, patch
import pandas as pd
from docx import Document


from document_processor import DocumentProcessor
from llm_summarizer import LLMSummarizer
from storage_manager import StorageManager

class TestDocumentProcessor(unittest.TestCase):
    
    def setUp(self):
        self.processor = DocumentProcessor()
    
    def test_clean_text(self):
        
        dirty_text = "  This   is   a   test   text   with   extra   spaces.  \n\n\n"
        cleaned = self.processor._clean_text(dirty_text)
        expected = "This is a test text with extra spaces."
        self.assertEqual(cleaned, expected)
    
    def test_txt_extraction(self):
   
       
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            f.write("This is a test document.\nIt has multiple lines.\n")
            temp_file = f.name
        
        try:
            
            mock_file = Mock()
            mock_file.name = "test.txt"
            mock_file.type = "text/plain"
            mock_file.read.return_value = "This is a test document.\nIt has multiple lines.\n".encode('utf-8')
            
            
            result = self.processor._extract_txt_text(mock_file)
            self.assertIn("This is a test document", result)
            self.assertIn("It has multiple lines", result)
            
        finally:
            os.unlink(temp_file)
    
    def test_csv_extraction(self):
       
        df = pd.DataFrame({
            'Name': ['Alice', 'Bob', 'Charlie'],
            'Age': [25, 30, 35],
            'City': ['NYC', 'LA', 'Chicago']
        })
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False) as f:
            df.to_csv(f.name, index=False)
            temp_file = f.name
        
        try:
           
            mock_file = Mock()
            mock_file.name = "test.csv"
            mock_file.type = "text/csv"
            
           
            with open(temp_file, 'rb') as f:
                mock_file.read.return_value = f.read()
            
           
            result = self.processor._extract_csv_text(mock_file)
            self.assertIn("Name", result)
            self.assertIn("Age", result)
            self.assertIn("City", result)
            
        finally:
            os.unlink(temp_file)
    
    def test_docx_extraction(self):
        
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph.')
        doc.add_paragraph('This is another paragraph.')
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            temp_file = f.name
        
        try:
           
            mock_file = Mock()
            mock_file.name = "test.docx"
            mock_file.type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
           
            with open(temp_file, 'rb') as f:
                mock_file.read.return_value = f.read()
            
            
            result = self.processor._extract_docx_text(mock_file)
            self.assertIn("Test Document", result)
            self.assertIn("test paragraph", result)
            
        finally:
            os.unlink(temp_file)

class TestLLMSummarizer(unittest.TestCase):
   
    def setUp(self):
        self.summarizer = LLMSummarizer()
    
    def test_available_models(self):
        
        models = self.summarizer.get_available_models()
        self.assertIn("gpt-3.5-turbo", models)
        self.assertIn("gpt-4", models)
        self.assertIn("gpt-4-turbo-preview", models)
    
    def test_calculate_max_tokens(self):
        
        tokens = self.summarizer._calculate_max_tokens(200)
        self.assertIsInstance(tokens, int)
        self.assertGreater(tokens, 0)
        self.assertLessEqual(tokens, 4000)
    
    def test_create_summary_prompt(self):
       
        text = "This is a test text for summarization."
        prompt = self.summarizer._create_summary_prompt(text, 100, 200)
        self.assertIn("comprehensive summary", prompt)
        self.assertIn("100", prompt)
        self.assertIn("200", prompt)
        self.assertIn(text, prompt)
    
    def test_estimate_cost(self):
       
        text = "This is a test text with some words for cost estimation."
        cost_info = self.summarizer.estimate_cost(text, "gpt-3.5-turbo")
        
        self.assertIn("input_tokens", cost_info)
        self.assertIn("output_tokens", cost_info)
        self.assertIn("total_tokens", cost_info)
        self.assertIn("estimated_cost_usd", cost_info)
        
        self.assertIsInstance(cost_info["estimated_cost_usd"], float)
        self.assertGreaterEqual(cost_info["estimated_cost_usd"], 0)

class TestStorageManager(unittest.TestCase):
    
    
    def setUp(self):
       
        self.temp_db = tempfile.NamedTemporaryFile(suffix='.db', delete=False)
        self.temp_db.close()
        
       
        self.storage = StorageManager()
        self.storage.db_file = self.temp_db.name
        self.storage._init_database()
    
    def tearDown(self):
        
        if os.path.exists(self.temp_db.name):
            os.unlink(self.temp_db.name)
        if os.path.exists(self.storage.json_file):
            os.unlink(self.storage.json_file)
        if os.path.exists(self.storage.csv_file):
            os.unlink(self.storage.csv_file)
    
    def test_save_and_retrieve_sqlite(self):
       
        summary_data = {
            'filename': 'test.pdf',
            'file_type': 'application/pdf',
            'file_size_kb': 100.5,
            'original_word_count': 500,
            'summary': 'This is a test summary.',
            'summary_word_count': 10,
            'model_used': 'gpt-3.5-turbo',
            'summary_length': 'Short (100-200 words)',
            'date': '2024-01-01T00:00:00',
            'extracted_text': 'This is the extracted text.'
        }
        
        
        self.storage.save_to_sqlite(summary_data)
        
        
        summaries = self.storage.get_all_summaries("sqlite")
        
        self.assertEqual(len(summaries), 1)
        self.assertEqual(summaries[0]['filename'], 'test.pdf')
        self.assertEqual(summaries[0]['summary'], 'This is a test summary.')
    
    def test_save_and_retrieve_json(self):
        
        summary_data = {
            'filename': 'test.txt',
            'file_type': 'text/plain',
            'file_size_kb': 50.0,
            'original_word_count': 200,
            'summary': 'This is a JSON test summary.',
            'summary_word_count': 8,
            'model_used': 'gpt-4',
            'summary_length': 'Medium (200-400 words)',
            'date': '2024-01-01T00:00:00',
            'extracted_text': 'This is the extracted text for JSON.'
        }
        
       
        self.storage.save_to_json(summary_data)
        
        
        summaries = self.storage.get_all_summaries("json")
        
        self.assertEqual(len(summaries), 1)
        self.assertEqual(summaries[0]['filename'], 'test.txt')
        self.assertEqual(summaries[0]['summary'], 'This is a JSON test summary.')
    
    def test_save_and_retrieve_csv(self):
        
        summary_data = {
            'filename': 'test.csv',
            'file_type': 'text/csv',
            'file_size_kb': 25.0,
            'original_word_count': 100,
            'summary': 'This is a CSV test summary.',
            'summary_word_count': 6,
            'model_used': 'gpt-3.5-turbo',
            'summary_length': 'Short (100-200 words)',
            'date': '2024-01-01T00:00:00',
            'extracted_text': 'This is the extracted text for CSV.'
        }
        
    
        self.storage.save_to_csv(summary_data)
        
     
        summaries = self.storage.get_all_summaries("csv")
        
        self.assertEqual(len(summaries), 1)
        self.assertEqual(summaries[0]['filename'], 'test.csv')
        self.assertEqual(summaries[0]['summary'], 'This is a CSV test summary.')
    
    def test_statistics(self):
        """Test statistics functionality"""
        
        summary_data = {
            'filename': 'test.pdf',
            'file_type': 'application/pdf',
            'file_size_kb': 100.0,
            'original_word_count': 500,
            'summary': 'Test summary.',
            'summary_word_count': 10,
            'model_used': 'gpt-3.5-turbo',
            'summary_length': 'Short (100-200 words)',
            'date': '2024-01-01T00:00:00',
            'extracted_text': 'Test text.'
        }
        
        self.storage.save_to_sqlite(summary_data)
        
       
        stats = self.storage.get_statistics()
        
        self.assertIsNotNone(stats)
        self.assertEqual(stats['total_documents'], 1)
        self.assertEqual(stats['total_words'], 500)
        self.assertEqual(stats['avg_summary_length'], 10.0)
        self.assertIn('gpt-3.5-turbo', stats['model_usage'])
    
    def test_search_functionality(self):
       
        
        summary_data = {
            'filename': 'research_paper.pdf',
            'file_type': 'application/pdf',
            'file_size_kb': 200.0,
            'original_word_count': 1000,
            'summary': 'This research paper discusses climate change impacts.',
            'summary_word_count': 15,
            'model_used': 'gpt-4',
            'summary_length': 'Medium (200-400 words)',
            'date': '2024-01-01T00:00:00',
            'extracted_text': 'Climate change research paper content.'
        }
        
        self.storage.save_to_sqlite(summary_data)
        
        
        results = self.storage.search_summaries("climate")
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]['filename'], 'research_paper.pdf')
        
       
        results = self.storage.search_summaries("research")
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]['filename'], 'research_paper.pdf')

def run_tests():
  
    print("Running Document Summarizer Tests...")
    print("=" * 50)
    
   
    test_suite = unittest.TestSuite()
    

    test_suite.addTest(unittest.makeSuite(TestDocumentProcessor))
    test_suite.addTest(unittest.makeSuite(TestLLMSummarizer))
    test_suite.addTest(unittest.makeSuite(TestStorageManager))
    
    
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(test_suite)
    
    
    print("=" * 50)
    if result.wasSuccessful():
        print(" All tests passed!")
        print(f" Tests run: {result.testsRun}")
        print(f"  Failures: {len(result.failures)}")
        print(f" Errors: {len(result.errors)}")
    else:
        print(" Some tests failed!")
        print(f" Tests run: {result.testsRun}")
        print(f" Failures: {len(result.failures)}")
        print(f" Errors: {len(result.errors)}")
        
        if result.failures:
            print("\n Failures:")
            for test, traceback in result.failures:
                print(f"  - {test}: {traceback}")
        
        if result.errors:
            print("\n Errors:")
            for test, traceback in result.errors:
                print(f"  - {test}: {traceback}")
    
    return result.wasSuccessful()

if __name__ == "__main__":
    success = run_tests()
    sys.exit(0 if success else 1) 